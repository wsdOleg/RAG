import csv
import re
from pathlib import Path

from app.config import Settings, getSettings
from app.services.ocrService import OcrService


class TextExtractionError(RuntimeError):
    pass


class TextExtractionService:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or getSettings()
        self.ocrService = OcrService(self.settings)

    def extractText(self, filePath: str) -> list[dict]:
        path = Path(filePath)
        suffix = path.suffix.lower()
        if suffix == ".txt":
            return [self.buildTextBlock(1, self.getTextContent(path), "text", "txt")]
        if suffix == ".csv":
            return [self.buildTextBlock(1, self.getCsvContent(path), "text", "csv")]
        if suffix == ".rtf":
            return [self.buildTextBlock(1, self.getRtfContent(path), "text", "rtf")]
        if suffix == ".pdf":
            return self.getPdfContent(path)
        if suffix == ".docx":
            return self.getDocxContent(path)
        if suffix == ".doc":
            raise TextExtractionError("DOC требует LibreOffice или antiword")
        if suffix in {".xlsx", ".xls"}:
            return [self.buildTextBlock(1, self.getSpreadsheetContent(path), "text", "spreadsheet")]
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
            result = self.ocrService.extractTextFromImage(str(path))
            result["pageNumber"] = 1
            if not result.get("text"):
                raise TextExtractionError("; ".join(result.get("warnings") or ["OCR не вернул текст"]))
            return [self.buildOcrBlock(result, "image_ocr")]
        raise TextExtractionError(f"Неподдерживаемый тип файла: {suffix}")

    def buildTextBlock(self, pageNumber: int | None, text: str, sourceType: str, extractionMethod: str, **extra: object) -> dict:
        return {
            "pageNumber": pageNumber,
            "text": self.postprocessExtractedText(text, sourceType, extractionMethod),
            "sourceType": sourceType,
            "extractionMethod": extractionMethod,
            **extra,
        }

    def buildOcrBlock(self, result: dict, extractionMethod: str, pageNumber: int | None = None) -> dict:
        return self.buildTextBlock(
            result.get("pageNumber") or pageNumber,
            result.get("text") or "",
            "ocr",
            extractionMethod,
            imageIndex=result.get("imageIndex"),
            imageName=result.get("imageName"),
            ocrConfidence=result.get("confidence"),
            ocrLanguage=self.settings.ocrLang,
            warnings=result.get("warnings") or [],
        )

    def getPdfContent(self, path: Path) -> list[dict]:
        try:
            from pypdf import PdfReader
        except Exception as error:
            raise TextExtractionError("Не найден pypdf") from error
        reader = PdfReader(str(path))
        blocks: list[dict] = []
        pagesForOcr: set[int] = set()
        for pageIndex, page in enumerate(reader.pages):
            pageText = page.extract_text() or ""
            pageNumber = pageIndex + 1
            if pageText.strip():
                blocks.append(self.buildTextBlock(pageNumber, pageText, "text", "pdf_text_layer"))
            if self.settings.ocrForceAllPdfPages or len(pageText.strip()) < self.settings.ocrMinTextLength:
                pagesForOcr.add(pageNumber)
        if pagesForOcr:
            for ocrBlock in self.ocrService.extractTextFromPdfPages(str(path)):
                pageNumber = ocrBlock.get("pageNumber")
                if pageNumber in pagesForOcr and ocrBlock.get("text"):
                    blocks.append(self.buildOcrBlock(ocrBlock, "pdf_ocr"))
                elif pageNumber in pagesForOcr and ocrBlock.get("warnings"):
                    blocks.append(self.buildTextBlock(pageNumber, "", "ocr", "pdf_ocr", warnings=ocrBlock.get("warnings") or []))
        return blocks

    def getDocxContent(self, path: Path) -> list[dict]:
        try:
            from docx import Document
        except Exception as error:
            raise TextExtractionError("Не найден python-docx") from error
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        tableLines: list[str] = []
        for tableIndex, table in enumerate(document.tables, start=1):
            tableLines.append(f"Таблица {tableIndex}")
            for rowIndex, row in enumerate(table.rows, start=1):
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    tableLines.append(f"Строка {rowIndex}: " + " | ".join(values))
        blocks: list[dict] = []
        fullText = "\n".join([*paragraphs, *tableLines]).strip()
        if fullText:
            blocks.append(self.buildTextBlock(1, fullText, "text", "docx_text"))
        for imageBlock in self.ocrService.extractImagesFromDocx(str(path)):
            if imageBlock.get("text"):
                blocks.append(self.buildOcrBlock(imageBlock, "docx_image_ocr", pageNumber=1))
            elif imageBlock.get("warnings"):
                blocks.append(self.buildTextBlock(1, "", "ocr", "docx_image_ocr", imageIndex=imageBlock.get("imageIndex"), imageName=imageBlock.get("imageName"), warnings=imageBlock.get("warnings") or []))
        return blocks

    def getTextContent(self, path: Path) -> str:
        for encoding in ("utf-8", "cp1251", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return path.read_text(encoding="utf-8", errors="ignore")

    def getCsvContent(self, path: Path) -> str:
        rawText = self.getTextContent(path)
        sample = rawText[:4096]
        try:
            dialect = csv.Sniffer().sniff(sample)
        except Exception:
            dialect = csv.excel
        rows: list[str] = []
        for rowIndex, row in enumerate(csv.reader(rawText.splitlines(), dialect)):
            if rowIndex >= 500:
                rows.append("... таблица обрезана после 500 строк")
                break
            rows.append(f"Строка {rowIndex + 1}: " + " | ".join(cell.strip() for cell in row))
        return "\n".join(rows)

    def getRtfContent(self, path: Path) -> str:
        rawText = self.getTextContent(path)
        try:
            from striprtf.striprtf import rtf_to_text
            return rtf_to_text(rawText)
        except Exception:
            cleanText = re.sub(r"\\'[0-9a-fA-F]{2}", " ", rawText)
            cleanText = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", cleanText)
            cleanText = cleanText.replace("{", " ").replace("}", " ")
            return re.sub(r"\s+", " ", cleanText).strip()

    def getSpreadsheetContent(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            try:
                from openpyxl import load_workbook
            except Exception as error:
                raise TextExtractionError("Не найден openpyxl") from error
            workbook = load_workbook(path, read_only=True, data_only=True)
            lines: list[str] = []
            for sheet in workbook.worksheets:
                lines.append(f"Лист: {sheet.title}")
                for rowIndex, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                    if rowIndex > 200:
                        lines.append("... лист обрезан после 200 строк")
                        break
                    values = ["" if value is None else str(value) for value in row]
                    if any(values):
                        lines.append(f"Строка {rowIndex}: " + " | ".join(values))
            return "\n".join(lines)
        try:
            import xlrd
        except Exception as error:
            raise TextExtractionError("Не найден xlrd") from error
        workbook = xlrd.open_workbook(str(path))
        lines: list[str] = []
        for sheet in workbook.sheets():
            lines.append(f"Лист: {sheet.name}")
            for rowIndex in range(min(sheet.nrows, 200)):
                values = [str(sheet.cell_value(rowIndex, columnIndex)) for columnIndex in range(sheet.ncols)]
                if any(values):
                    lines.append(f"Строка {rowIndex + 1}: " + " | ".join(values))
        return "\n".join(lines)

    def postprocessExtractedText(self, text: str, sourceType: str, extractionMethod: str) -> str:
        normalizedText = (text or "").replace("\r\n", "\n").replace("\r", "\n")
        if sourceType == "ocr" or extractionMethod in {"pdf_text_layer", "docx_text"}:
            normalizedText = self.removeServiceLines(normalizedText)
            normalizedText = self.normalizeLegalText(normalizedText)
        return normalizedText.strip()

    def removeServiceLines(self, text: str) -> str:
        keptLines: list[str] = []
        for rawLine in text.split("\n"):
            line = rawLine.strip()
            if not line:
                keptLines.append("")
                continue
            lowerLine = line.lower()
            if "передан через диадок" in lowerLine:
                continue
            if re.fullmatch(r"страница\s+\d+\s+из\s+\d+", lowerLine):
                continue
            if re.fullmatch(r"[\dA-Za-zА-Яа-я+\-=/|]{10,}", line) and not re.search(r"[А-Яа-яЁё]{3,}", line):
                continue
            if re.fullmatch(r"[_\-/\\=|`~^]+", line):
                continue
            if len(line) <= 2 and not re.search(r"\d", line):
                continue
            keptLines.append(line)
        return "\n".join(keptLines)

    def normalizeLegalText(self, text: str) -> str:
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
        paragraphs: list[str] = []
        currentParagraph = ""

        for line in lines:
            if not line:
                if currentParagraph:
                    paragraphs.append(self.normalizeLegalLine(currentParagraph))
                    currentParagraph = ""
                continue

            if self.isStandaloneLegalLine(line):
                if currentParagraph:
                    paragraphs.append(self.normalizeLegalLine(currentParagraph))
                    currentParagraph = ""
                paragraphs.append(self.normalizeLegalLine(line))
                continue

            if not currentParagraph:
                currentParagraph = line
                continue

            if self.shouldMergeLegalLines(currentParagraph, line):
                currentParagraph = f"{currentParagraph} {line.lstrip('-–— ')}"
            else:
                paragraphs.append(self.normalizeLegalLine(currentParagraph))
                currentParagraph = line

        if currentParagraph:
            paragraphs.append(self.normalizeLegalLine(currentParagraph))

        compactParagraphs: list[str] = []
        for paragraph in paragraphs:
            cleanedParagraph = paragraph.strip()
            if not cleanedParagraph:
                continue
            compactParagraphs.append(cleanedParagraph)
        return "\n\n".join(compactParagraphs)

    def normalizeLegalLine(self, text: str) -> str:
        normalizedText = re.sub(r"\s+", " ", text).strip()
        replacements = {
            "Ne ": "№ ",
            " Nе ": " № ",
            " Nе": " №",
            "No ": "№ ",
            "ct. ": "ст. ",
            "cr. ": "ст. ",
            "r. ": "г. ",
            "66 этом": "об этом",
            "66 этом ": "об этом ",
            "Saa5": "SaaS",
        }
        for sourceText, targetText in replacements.items():
            normalizedText = normalizedText.replace(sourceText, targetText)
        normalizedText = re.sub(r"\bN[eoо]\s*(\d)", r"№ \1", normalizedText)
        normalizedText = re.sub(r"\bст\.\s*(\d)", r"ст. \1", normalizedText)
        normalizedText = normalizedText.replace(" ,", ",").replace(" .", ".").replace(" :", ":").replace(" ;", ";")
        normalizedText = re.sub(r"([А-Яа-яЁё])\s{2,}([А-Яа-яЁё])", r"\1 \2", normalizedText)
        normalizedText = re.sub(r"\s+([,.;:])", r"\1", normalizedText)
        return normalizedText.strip()

    def isStandaloneLegalLine(self, line: str) -> bool:
        if re.match(r"^\d+(\.\d+)*\.?$", line):
            return True
        if re.match(r"^\d+(\.\d+)+\s", line):
            return True
        if re.match(r"^(договор|соглашение|акт|приложение)\b", line.lower()):
            return True
        if re.match(r"^[А-ЯЁA-Z][А-ЯЁA-Z0-9 №«»\"().,-]{8,}$", line):
            return True
        return False

    def shouldMergeLegalLines(self, currentLine: str, nextLine: str) -> bool:
        if self.isStandaloneLegalLine(nextLine):
            return False
        if re.match(r"^\d+(\.\d+)+\s", nextLine):
            return False
        if currentLine.endswith((".", "!", "?", ";", ":")):
            return False
        return True
